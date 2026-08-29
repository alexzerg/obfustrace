from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUTPUT = Path("templates/doctavian/emergency-travel-request.docx")


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str = "0F172A") -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_label_value(document: Document, label: str, expression: str) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.75)
    table.columns[1].width = Inches(4.8)
    set_cell_text(table.cell(0, 0), label.upper(), bold=True, color="64748B")
    set_cell_text(table.cell(0, 1), expression, bold=True)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.7)
section.right_margin = Inches(0.7)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(10)
styles["Title"].font.name = "Arial"
styles["Title"].font.size = Pt(24)
styles["Title"].font.bold = True
styles["Heading 1"].font.name = "Arial"
styles["Heading 1"].font.size = Pt(15)
styles["Heading 1"].font.color.rgb = RGBColor(15, 118, 110)

banner = doc.add_table(rows=1, cols=2)
banner.alignment = WD_TABLE_ALIGNMENT.CENTER
banner.autofit = False
banner.columns[0].width = Inches(4.9)
banner.columns[1].width = Inches(1.65)
shade(banner.cell(0, 0), "0F172A")
shade(banner.cell(0, 1), "0D9488")
set_cell_text(banner.cell(0, 0), "MICRO-EMBASSY", bold=True, color="FFFFFF")
set_cell_text(banner.cell(0, 1), "TEMPORARY", bold=True, color="FFFFFF")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("EMERGENCY TRAVEL DOCUMENT REQUEST")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(15, 23, 42)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Synthetic hackathon demonstration — not an official government form")
run.bold = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(190, 18, 60)

add_label_value(doc, "Case", "{!Case[0].CaseId}")
add_label_value(doc, "Generated", "{!Case[0].GeneratedAt}")
add_label_value(doc, "Recipient", "{!Case[0].Recipient.Organization}")
add_label_value(doc, "Purpose", "{!Case[0].Recipient.Purpose}")

for heading in ["Traveler identity", "Incident details", "Confirmed travel"]:
    doc.add_heading(heading, level=1)
    if heading == "Traveler identity":
        add_label_value(doc, "Full name", "{!Case[0].Traveler.FullName}")
        add_label_value(doc, "Nationality", "{!Case[0].Traveler.Nationality}")
        add_label_value(doc, "Date of birth", "{!Case[0].Traveler.DateOfBirth}")
        add_label_value(doc, "Document no.", "{!Case[0].Traveler.DocumentNumber}")
    elif heading == "Incident details":
        add_label_value(doc, "Incident", "{!Case[0].Incident.Type}")
        add_label_value(doc, "Date", "{!Case[0].Incident.Date}")
        add_label_value(doc, "Location", "{!Case[0].Incident.Location}")
        add_label_value(doc, "Police ref.", "{!Case[0].Incident.PoliceReportReference}")
        paragraph = doc.add_paragraph("{!Case[0].Incident.Statement}")
        paragraph.paragraph_format.space_before = Pt(6)
    else:
        add_label_value(doc, "Route", "{!Case[0].Travel.Route}")
        add_label_value(doc, "Flight", "{!Case[0].Travel.FlightNumber}")
        add_label_value(doc, "Departure", "{!Case[0].Travel.Departure}")
        add_label_value(doc, "Booking ref.", "{!Case[0].Travel.BookingReference}")


doc.add_heading("Evidence supplied", level=1)
doc.add_paragraph('<mdoc:table name="evidenceTable" value="{!Case[0].Evidence}" variable="evidence" repeatingRow="1">')
evidence_table = doc.add_table(rows=2, cols=3)
evidence_table.style = "Table Grid"
evidence_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, label in enumerate(["Evidence", "Reference", "Review status"]):
    shade(evidence_table.cell(0, i), "0F172A")
    set_cell_text(evidence_table.cell(0, i), label, bold=True, color="FFFFFF")
for i, expression in enumerate([
    "{!#evidence#.Type}",
    "{!#evidence#.Reference}",
    "{!#evidence#.Status}",
]):
    set_cell_text(evidence_table.cell(1, i), expression)
doc.add_paragraph('</mdoc:table name="evidenceTable">')


doc.add_heading("Traveler declaration", level=1)
for expression in [
    "{!Case[0].Declaration.Accuracy}",
    "{!Case[0].Declaration.Consent}",
    "{!Case[0].Declaration.Retention}",
]:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run(expression)

signature = doc.add_table(rows=2, cols=2)
signature.alignment = WD_TABLE_ALIGNMENT.CENTER
signature.autofit = False
signature.columns[0].width = Inches(3.2)
signature.columns[1].width = Inches(3.2)
set_cell_text(signature.cell(0, 0), "Traveler signature: __________________________")
set_cell_text(signature.cell(0, 1), "Date: __________________________")
set_cell_text(signature.cell(1, 0), "Consular review: __________________________")
set_cell_text(signature.cell(1, 1), "Case status: __________________________")

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer.add_run("Micro-Embassy · Purpose-bound emergency package · Case {!Case[0].CaseId}")
footer_run.font.size = Pt(8)
footer_run.font.color.rgb = RGBColor(100, 116, 139)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
